"""Testes para core/doc_converter.py."""
from pathlib import Path

import pytest

from core.doc_converter import EXTENSOES_DOC, _decodificar_textutil, doc_to_md
from core.utils import ModoImagem


def _criar_docx_simples(path: Path, texto: str = "Conteúdo de teste Word.") -> None:
    """Cria um .docx mínimo válido com texto."""
    from docx import Document
    doc = Document()
    doc.add_paragraph(texto)
    doc.save(str(path))


def _criar_docx_estruturado(path: Path) -> None:
    """Cria .docx com título, parágrafos e lista."""
    from docx import Document
    doc = Document()
    doc.add_heading("Relatório de Testes", level=1)
    doc.add_paragraph("Primeiro parágrafo do documento.")
    doc.add_heading("Seção 2", level=2)
    doc.add_paragraph("Segundo parágrafo.")
    doc.save(str(path))


def test_docx_texto_simples(tmp_path):
    """docx com texto → MD com conteúdo."""
    path = tmp_path / "simples.docx"
    _criar_docx_simples(path, "Texto simples para teste.")
    md = doc_to_md(path)
    assert len(md) > 5
    assert "Texto simples" in md or len(md) > 0


def test_docx_estruturado(tmp_path):
    """docx com headers → MD com conteúdo."""
    path = tmp_path / "estruturado.docx"
    _criar_docx_estruturado(path)
    md = doc_to_md(path)
    assert len(md) > 20
    # mammoth preserva headers como # / ##
    assert "Relatório" in md or "Seção" in md or len(md) > 30


def test_docx_nao_existe():
    """FileNotFoundError para arquivo inexistente."""
    with pytest.raises(FileNotFoundError):
        doc_to_md(Path("/nao/existe.docx"))


def test_extensao_invalida(tmp_path):
    """ValueError para extensão não suportada."""
    path = tmp_path / "arquivo.txt"
    path.write_text("conteúdo", encoding="utf-8")
    with pytest.raises(ValueError, match="Extensão não suportada"):
        doc_to_md(path)


def test_docx_corrompido(tmp_path):
    """RuntimeError para .docx corrompido."""
    path = tmp_path / "corrompido.docx"
    path.write_text("isto não é um docx", encoding="utf-8")
    with pytest.raises(RuntimeError):
        doc_to_md(path)


def test_extensoes_doc_set():
    """.doc e .docx estão em EXTENSOES_DOC."""
    assert ".doc" in EXTENSOES_DOC
    assert ".docx" in EXTENSOES_DOC


def test_decodificar_textutil_latin1():
    """Bytes Latin-1 com acentos PT-BR → replacement chars (não falha).

    textutil emite UTF-8 nativamente. Se bytes Latin-1 vazarem da fonte
    original, errors='replace' produz U+FFFD em vez de decodificar como
    Latin-1 — isso é aceitável porque o caso não ocorre na prática
    (textutil sempre converte para UTF-8).
    """
    # 'ção' em Latin-1/CP1252: 0xE7=ç 0xE3=ã 0x6F=o — bytes inválidos em UTF-8
    resultado = _decodificar_textutil(b"\xe7\xe3o")
    assert isinstance(resultado, str)
    assert "o" in resultado  # 'o' (0x6F) é válido em qualquer encoding
    assert "\ufffd" in resultado  # bytes inválidos viram replacement char


def test_decodificar_textutil_utf8():
    """Saída já em UTF-8 (padrão do textutil) é decodificada corretamente."""
    assert _decodificar_textutil("olá ção".encode()) == "olá ção"


def test_decodificar_textutil_nunca_levanta():
    """Qualquer sequência de bytes decodifica sem exceção (errors='replace')."""
    # 0x81/0x8D/0x90 são indefinidos em CP1252 mas válidos em Latin-1
    resultado = _decodificar_textutil(b"\x81\x8d\x90\xff")
    assert isinstance(resultado, str)


def test_doc_ole_real_via_textutil(tmp_path):
    """.doc OLE binário real (gerado pelo próprio textutil) → conversão end-to-end.

    Cobre o caminho real de _doc_para_md com subprocess de verdade (sem mock),
    validando que a troca antiword→textutil funciona com um arquivo OLE
    legítimo (assinatura D0 CF 11 E0), não apenas com .docx disfarçado de .doc.
    """
    import subprocess as sp
    txt_path = tmp_path / "fonte.txt"
    txt_path.write_text("Conteúdo de teste com acentuação: ção ã é.", encoding="utf-8")
    doc_path = tmp_path / "legado.doc"
    sp.run(
        ["/usr/bin/textutil", "-convert", "doc", str(txt_path), "-output", str(doc_path)],
        check=True,
        timeout=30,
    )
    with open(doc_path, "rb") as f:
        assert f.read(4) == b"\xd0\xcf\x11\xe0"  # confirma OLE binário real

    md = doc_to_md(doc_path)
    assert "Conteúdo de teste" in md
    assert "ção" in md


def test_batch_docx(tmp_path):
    """batch_convert processa .docx corretamente."""
    from core.batch import StatusArquivo, batch_convert
    path = tmp_path / "entrada" / "doc.docx"
    path.parent.mkdir()
    _criar_docx_simples(path)
    resultados = batch_convert(
        origem=path,
        destino=tmp_path / "saida",
        workers=1,
    )
    assert len(resultados) == 1
    assert resultados[0].status == StatusArquivo.CONCLUIDO
    assert resultados[0].destino is not None
    assert resultados[0].destino.exists()


def test_batch_doc_extensao_reconhecida(tmp_path):
    """batch_convert converte .doc que é na verdade um .docx (zip "PK").

    Determinístico: o sniff de assinatura roteia "PK" para mammoth, sem
    depender de antiword. Antes o teste só asseria `!= IGNORADO`, passando
    também em ERRO — não validava conversão real.
    """
    from core.batch import StatusArquivo, batch_convert
    path = tmp_path / "entrada" / "doc.doc"
    path.parent.mkdir()
    _criar_docx_simples(path, "Documento Word salvo como .doc")
    resultados = batch_convert(
        origem=path,
        destino=tmp_path / "saida",
        workers=1,
    )
    assert len(resultados) == 1
    assert resultados[0].status == StatusArquivo.CONCLUIDO
    assert resultados[0].destino is not None
    assert resultados[0].destino.exists()
    assert resultados[0].destino.read_text(encoding="utf-8").strip() != ""


# ── T19: unificação da política de imagens com o PDF ─────────────────────────

def _criar_docx_com_imagem(path: Path, qtd: int = 1, cor=(200, 50, 50)) -> None:
    """Cria .docx com texto antes/depois e `qtd` imagens embutidas."""
    import tempfile

    from docx import Document
    from PIL import Image

    doc = Document()
    doc.add_paragraph("Antes da imagem:")
    for i in range(qtd):
        p = doc.add_paragraph()
        run = p.add_run()
        img = Image.new("RGB", (30, 30), color=((cor[0] + i * 40) % 255, cor[1], cor[2]))
        with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
            img.save(tmp.name)
            run.add_picture(tmp.name)
    doc.add_paragraph("Depois da imagem.")
    doc.save(str(path))


def test_docx_default_descarta_base64(tmp_path):
    """Default (transcrever): sem base64 e sem assets — comportamento novo.

    O mammoth embutia data-URI base64 por padrão (rejeitado na D2); o T19
    unifica com o PDF: default descarta a imagem.
    """
    path = tmp_path / "com_imagem.docx"
    _criar_docx_com_imagem(path)

    md = doc_to_md(path)

    assert "base64" not in md
    assert "![" not in md
    assert "Antes da imagem" in md and "Depois da imagem" in md
    assert not (tmp_path / "com_imagem_assets").exists()


def test_docx_extrair_posicao_preservada(tmp_path):
    """extrair: asset no disco + link no ponto exato do texto (posição!)."""
    path = tmp_path / "com_imagem.docx"
    _criar_docx_com_imagem(path)
    saida = tmp_path / "out"
    assets = saida / "com_imagem_assets"
    saida.mkdir()

    md = doc_to_md(
        path, modo_imagem=ModoImagem.extrair, assets_dir=assets, md_dir=saida
    )

    assert (assets / "img_0001.png").exists()
    # Posição: entre "Antes da imagem" e "Depois da imagem" (≠ PDF, que anexa no fim)
    pos_antes = md.index("Antes da imagem")
    pos_link = md.index("![](com_imagem_assets/img_0001.png)")
    pos_depois = md.index("Depois da imagem")
    assert pos_antes < pos_link < pos_depois


def test_docx_ambos_ocr_como_alt(tmp_path):
    """ambos: OCR da imagem vira alt-text do link."""
    from unittest.mock import patch

    path = tmp_path / "com_imagem.docx"
    _criar_docx_com_imagem(path)

    with patch("core.image_converter.image_to_md", return_value="TEXTO DA IMAGEM OCR"):
        md = doc_to_md(path, modo_imagem=ModoImagem.ambos)

    assert "![TEXTO DA IMAGEM OCR](com_imagem_assets/img_0001.png)" in md


def test_docx_ignorar_descarta(tmp_path):
    """ignorar: imagem descartada sem escrever nada."""
    path = tmp_path / "com_imagem.docx"
    _criar_docx_com_imagem(path)

    md = doc_to_md(path, modo_imagem=ModoImagem.ignorar)

    assert "![" not in md
    assert "base64" not in md
    assert not (tmp_path / "com_imagem_assets").exists()


def test_docx_obsidian_wikilink(tmp_path):
    """obsidian (D1): wikilink ![[...]] no MD."""
    path = tmp_path / "com_imagem.docx"
    _criar_docx_com_imagem(path)
    vault = tmp_path / "vault"
    (vault / "attachments").mkdir(parents=True)

    md = doc_to_md(
        path,
        modo_imagem=ModoImagem.extrair,
        assets_dir=vault / "attachments",
        md_dir=vault,
        wikilinks=True,
        prefixo_nome="com_imagem__",
    )

    assert "![[com_imagem__img_0001.png]]" in md
    assert (vault / "attachments" / "com_imagem__img_0001.png").exists()


def test_docx_dedup_mesma_imagem(tmp_path):
    """Duas imagens idênticas → 1 arquivo, 2 links (D4)."""
    from docx import Document

    path = tmp_path / "duplicada.docx"
    doc = Document()
    doc.add_paragraph("Antes:")
    for _ in range(2):
        p = doc.add_paragraph()
        run = p.add_run()
        import tempfile

        from PIL import Image
        img = Image.new("RGB", (30, 30), color=(10, 200, 10))
        with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
            img.save(tmp.name)
            run.add_picture(tmp.name)
    doc.add_paragraph("Depois.")
    doc.save(str(path))

    md = doc_to_md(path, modo_imagem=ModoImagem.extrair)

    assert md.count("![](") == 2          # 2 links
    assets_dir = tmp_path / "duplicada_assets"
    assert len(list(assets_dir.iterdir())) == 1  # 1 arquivo no disco


def test_docx_limite_imagens_por_documento(tmp_path):
    """Limite por documento → descarta o restante com aviso (gate Sentinel)."""
    from unittest.mock import patch

    path = tmp_path / "muitas.docx"
    _criar_docx_com_imagem(path, qtd=4)
    avisos: list[str] = []

    with patch("core.image_assets._MAX_IMAGENS_PDF", 2):
        md = doc_to_md(path, modo_imagem=ModoImagem.extrair, avisos=avisos)

    assert md.count("![](") == 2
    assert any("limite" in a for a in avisos)
