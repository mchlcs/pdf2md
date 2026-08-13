"""Testes para core/image_assets.py — FIX 1 (prefixo ASCII seguro).

Regressão: o prefixo derivado do stem do arquivo (ex: "Relatório Final__"
com espaço/acento) violava o regex de nome seguro `^[A-Za-z0-9._-]+$` em
--obsidian/--assets-dir → ValueError em `caminho_seguro` → o arquivo
inteiro virava ERRO no batch.
"""
import re
from pathlib import Path

from core.image_assets import slugificar_ascii
from core.utils import ModoImagem


def test_slugificar_ascii_remove_acentos_e_espacos():
    """'Relatório Final__' → 'Relatorio_Final__' (FIX 1)."""
    assert slugificar_ascii("Relatório Final__") == "Relatorio_Final__"


def test_slugificar_ascii_preserva_hifen_e_ponto():
    """Hífen e ponto (válidos no regex de nome seguro) são preservados."""
    assert slugificar_ascii("rel-pdf__") == "rel-pdf__"
    assert slugificar_ascii("v1.2__") == "v1.2__"


def test_slugificar_ascii_sempre_ascii_seguro():
    """Qualquer entrada → casa o regex de nome seguro (sem ValueError)."""
    padrao = re.compile(r"^[A-Za-z0-9._-]+$")
    for entrada in [
        "Árvore de Decisão__",
        "ação!@# ção__",
        "Plano B (2026)__",
        "nota__",
        "  espaços  __",
    ]:
        assert padrao.match(slugificar_ascii(entrada)), entrada


def test_resolver_assets_prefixo_saneado_assets_dir():
    """--assets-dir com stem acentuado → prefixo ASCII (batch._resolver_assets)."""
    from core.batch import _resolver_assets

    origem = Path("/tmp/Relatório Final.pdf")
    destino = Path("/tmp/saida/Relatório Final.md")
    assets, prefixo = _resolver_assets(origem, destino, obsidian=False, assets_dir_str="/tmp/assets")
    assert prefixo == "Relatorio_Final__"


def test_resolver_assets_prefixo_saneado_obsidian():
    """--obsidian com stem acentuado → prefixo ASCII."""
    from core.batch import _resolver_assets

    origem = Path("/tmp/Plano de Ação.pdf")
    destino = Path("/tmp/vault/Plano de Ação.md")
    assets, prefixo = _resolver_assets(origem, destino, obsidian=True, assets_dir_str=None)
    assert prefixo == "Plano_de_Acao__"
    assert str(assets).endswith("attachments")


def test_resolver_assets_default_sem_prefixo():
    """Sem diretório compartilhado: prefixo vazio (comportamento intacto)."""
    from core.batch import _resolver_assets

    origem = Path("/tmp/Relatório Final.pdf")
    destino = Path("/tmp/saida/Relatório Final.md")
    assets, prefixo = _resolver_assets(origem, destino, obsidian=False, assets_dir_str=None)
    assert prefixo == ""
    assert assets.name == "Relatório Final_assets"


def test_batch_assets_dir_stem_acentuado_nao_vira_erro(tmp_path):
    """FIX 1 E2E: docx 'Relatório Final.docx' + --assets-dir → CONCLUIDO.

    Regressão: prefixo 'Relatório Final__' violava o regex de nome seguro
    → ValueError em caminho_seguro → arquivo inteiro virava ERRO.
    """
    import tempfile

    from docx import Document
    from PIL import Image

    from core.batch import StatusArquivo, batch_convert

    path = tmp_path / "Relatório Final.docx"
    doc = Document()
    doc.add_paragraph("Nota com figura.")
    p = doc.add_paragraph()
    run = p.add_run()
    img = Image.new("RGB", (30, 30), color=(10, 200, 90))
    with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
        img.save(tmp.name)
        run.add_picture(tmp.name)
    doc.save(str(path))

    assets_globais = tmp_path / "assets_globais"
    resultados = batch_convert(
        origem=path,
        destino=tmp_path / "saida",
        workers=1,
        modo_imagem=ModoImagem.extrair,
        assets_dir=assets_globais,
    )

    assert resultados[0].status == StatusArquivo.CONCLUIDO
    assert (assets_globais / "Relatorio_Final__img_0001.png").exists()
